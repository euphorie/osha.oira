from .. import _
from ..interfaces import IOSHAContentSkinLayer
from datetime import date
from docx.api import Document
from euphorie.client.docx.compiler import _sanitize_html
from euphorie.client.docx.compiler import delete_paragraph
from euphorie.client.docx.compiler import IdentificationReportCompiler
from euphorie.client.docx.html import HtmlToWord
from euphorie.client.docx.views import IdentificationReportDocxView
from euphorie.content import survey
from euphorie.content.interfaces import IQuestionContainer
from euphorie.content.profilequestion import IProfileQuestion
from euphorie.content.risk import IRisk
from euphorie.content.solution import ISolution
from euphorie.content.survey import get_tool_type
from euphorie.content.utils import IToolTypesInfo
from five import grok
from plone.app.dexterity.behaviors.metadata import DCFieldProperty
from plone.app.dexterity.behaviors.metadata import MetadataBase
from plone.autoform.interfaces import IFormFieldProvider
from plone.directives import dexterity, form
from plone.namedfile import field as filefield
from plonetheme.nuplone.utils import formatDate
from plonetheme.nuplone.z3cform.directives import depends
from zope import interface
from zope import schema
from zope.component import getMultiAdapter
from zope.component import getUtility
from zope.i18n import translate
import z3c.form


grok.templatedir("templates")

help_default_probability = _(
    u"help_default_probability",
    default=u"Indicate how likely occurence of this risk is in a normal situation.")
help_default_frequency = _(
    u"help_default_frequency",
    default=u"Indicate how often this risk occurs in a normal situation.")
help_default_severity =_(
    u"help_default_severity",
    default=u"Indicate the severity if this risk occurs.")


class IOSHASurvey(form.Schema):
    """ Adds a logo, URL and name of an external reference site to a survey """
    enable_external_site_link = schema.Bool(
        title=_("label_external_site_enabled",
                default=u"Include a logo which links to an external "
                        u"website."),
        description=_("help_external_site_enabled",
                      default=u"Tick this option if you would like to create "
                      u"a hyperlink on the OiRA tool which points to an "
                      u"external website. The hyperlink will be in the form "
                      u"of a logo image."),
        required=False,
        default=False)

    depends("IOSHASurvey.external_site_url",
            "IOSHASurvey.enable_external_site_link",
            "on")
    external_site_url = schema.URI(
        title=_("label_external_site_url", default=u"External site URL"),
        description=_("help__external_site_url",
                      default=u"This is the URL of an external site that is "
                      u"linked to. Clicking the logo or the name will take "
                      u"the user to this URL."),
        required=False)

    depends("IOSHASurvey.external_site_name",
            "IOSHASurvey.enable_external_site_link",
            "on")
    external_site_name = schema.TextLine(
        title=_("label_external_site_name", default=u"External site name"),
        description=_("help_external_site_name",
                      default=u"This is the name of the external site that is "
                      u"linked to. It will appear next to the logo."),
        required=False)

    depends("IOSHASurvey.external_site_logo",
            "IOSHASurvey.enable_external_site_link",
            "on")
    external_site_logo = filefield.NamedBlobImage(
        title=_("label_external_site_logo", default=u"External site logo"),
        description=_(
            "help_image_upload",
            default=u"Upload an image. Make sure your image is of format "
                    u"png, jpg or gif and does not contain any special "
                    u"characters."),
        required=False)

    enable_custom_evaluation_descriptions = schema.Bool(
        title=_("label_enable_custom_evaluation_descriptions",
                default=u"The criteria applied to evaluate risks are specific "
                u"of this tool? (If not, the common criteria descriptions "
                u"will apply)."),
        description=_("help_enable_custom_evaluation_descriptions",
                      default=u"Tick this option if you would like to define "
                      u"your own descriptions for the criteria of the "
                      u"evaluation algorithm. The user will see them as hints "
                      u"when answering the questions to calculate the "
                      u"priority of a risk."),
        required=False,
        default=False)

    depends("IOSHASurvey.description_probability",
            "IOSHASurvey.enable_custom_evaluation_descriptions",
            "on")
    description_probability = schema.Text(
        title=_(u"Probability"),
        description=_(
            u"description_criteria_explanation",
            default=u'Provide your custom explanation here, to override this '
                    u'default explanation: "${default_explanation}"',
            mapping={
                u'default_explanation': help_default_probability}
        ),
        required=False,
    )

    depends("IOSHASurvey.description_frequency",
            "IOSHASurvey.enable_custom_evaluation_descriptions",
            "on")
    description_frequency = schema.Text(
        title=_(u"Frequency"),
        description=_(
            u"description_criteria_explanation",
            default=u'Provide your custom explanation here, to override this '
                    u'default explanation: "${default_explanation}"',
            mapping={
                u'default_explanation': help_default_frequency}
        ),
        required=False,
    )

    depends("IOSHASurvey.description_severity",
            "IOSHASurvey.enable_custom_evaluation_descriptions",
            "on")
    description_severity = schema.Text(
        title=_(u"Severity"),
        description=_(
            u"description_criteria_explanation",
            default=u'Provide your custom explanation here, to override this '
                    u'default explanation: "${default_explanation}"',
            mapping={
                u'default_explanation': help_default_severity}
        ),
        required=False,
    )

interface.alsoProvides(IOSHASurvey, IFormFieldProvider)


class IOSHASurveyMarker(survey.ISurvey):
    """ Marker interface so that we can register more specific adapters for
        OSHA's survey object.
    """

interface.classImplements(survey.Survey, IOSHASurveyMarker)


class OSHASurvey(MetadataBase):
    enable_external_site_link = \
        DCFieldProperty(IOSHASurvey['enable_external_site_link'])
    external_site_url = DCFieldProperty(IOSHASurvey['external_site_url'])
    external_site_name = DCFieldProperty(IOSHASurvey['external_site_name'])
    external_site_logo = DCFieldProperty(IOSHASurvey['external_site_logo'])
    enable_custom_evaluation_descriptions = DCFieldProperty(IOSHASurvey['enable_custom_evaluation_descriptions'])
    description_probability = DCFieldProperty(IOSHASurvey['description_probability'])
    description_frequency = DCFieldProperty(IOSHASurvey['description_frequency'])
    description_severity = DCFieldProperty(IOSHASurvey['description_severity'])


class OSHASurveyEditForm(dexterity.EditForm):
    grok.context(survey.ISurvey)
    grok.layer(IOSHAContentSkinLayer)

    def updateWidgets(self):
        result = super(OSHASurveyEditForm, self).updateWidgets()
        evaluation_optional = self.widgets.get('evaluation_optional')
        evaluation_optional.mode = z3c.form.interfaces.HIDDEN_MODE
        if self.context.aq_parent.evaluation_algorithm == 'french':
            description_probability = self.widgets.get('IOSHASurvey.description_probability')
            description_probability.mode = z3c.form.interfaces.HIDDEN_MODE
        return result


class OSHASurveyView(survey.View):
    grok.layer(IOSHAContentSkinLayer)
    grok.template("survey_view")

    def modules_and_profile_questions(self):
        return [self._morph(child) for child in self.context.values()]

    def _morph(self, child):
        state = getMultiAdapter(
                    (child, self.request),
                    name="plone_context_state")

        return dict(id=child.id,
                    title=child.title,
                    url=state.view_url(),
                    is_profile_question=IProfileQuestion.providedBy(child))


class ContentsOfSurveyCompiler(IdentificationReportCompiler):
    def __init__(self, context, request=None):
        """ Read the docx template and initialize some instance attributes
        that will be used to compile the template
        """
        self.context = context
        self.request = request
        self.template = Document(self._template_filename)
        self.use_existing_measures = False
        self.tool_type = get_tool_type(self.context)
        self.tti = getUtility(IToolTypesInfo)
        self.italy_special = False

    def set_session_title_row(self, data):

        request = self.request

        # Remove existing paragraphs
        for paragraph in self.template.paragraphs:
            delete_paragraph(paragraph)

        header = self.template.sections[0].header
        header_table = header.tables[0]
        header_table.cell(0, 0).paragraphs[0].text = data["title"]
        header_table.cell(0, 1).paragraphs[0].text = formatDate(request, date.today())

        if getattr(self.context, "published"):
            footer_txt = u"This OiRA tool was last published {date}.".format(
                date=self.context.published.strftime("%Y/%m/%d")
            )
        else:
            footer_txt = u"This OiRA tool is currently not published."

        footer = self.template.sections[0].footer
        paragraph = footer.tables[0].cell(0, 0).paragraphs[0]
        paragraph.style = "Footer"
        paragraph.text = footer_txt

    def add_report_section(self, nodes, heading, **extra):
        doc = self.template
        doc.add_paragraph(heading, style="Heading 1")

        for node in nodes:
            title = u"[{0}] {1}".format(
                translate(_(node.typus), target_language=self.lang), node.title
            )
            number = node.number

            doc.add_paragraph(
                u"%s %s" % (number, title), style="Heading %d" % (node.depth + 1)
            )

            if node.typus == "Risk":
                doc.add_paragraph(
                    u"[%s] %s"
                    % (
                        translate(
                            _(
                                "label_problem_description",
                                default=u"Negative statement",
                            ),
                            target_language=self.lang,
                        ),
                        node.problem_description,
                    ),
                    style="Measure Heading",
                )

            description = node.description

            doc = HtmlToWord(_sanitize_html(description or ""), doc)

            if node.typus != "Risk":
                continue

            if not extra.get("skip_legal_references", False):
                legal_reference = getattr(node, "legal_reference", None)
                if legal_reference and legal_reference.strip():
                    doc.add_paragraph()
                    legal_heading = translate(
                        _(
                            "header_legal_references",
                            default=u"Legal and policy references",
                        ),
                        target_language=self.lang,
                    )
                    doc.add_paragraph(legal_heading, style="Legal Heading")
                    doc = HtmlToWord(_sanitize_html(legal_reference), doc)


class Node(object):

    title = ""
    typus = ""
    depth = 0
    number = ""
    description = ""
    legal_reference = None
    problem_description = None

    def __init__(
        self,
        title,
        typus,
        depth,
        number,
        description="",
        legal_reference=None,
        problem_description=None,
    ):
        self.title = title
        self.typus = typus
        self.depth = depth
        self.number = number
        self.description = description
        self.legal_reference = legal_reference
        self.problem_description = problem_description


class ContentsOfSurvey(IdentificationReportDocxView):

    _compiler = ContentsOfSurveyCompiler
    nodes = []

    def __init__(self, request, context):
        super(ContentsOfSurvey, self).__init__(request, context)
        self.nodes = []

    def AddToTree(self, node, depth, number):
        legal_reference = None
        problem_description = None
        if IRisk.providedBy(node):
            typus = "Risk"
            legal_reference = getattr(node, "legal_reference", None)
            problem_description = getattr(node, "problem_description", None)
        elif IProfileQuestion.providedBy(node):
            typus = "Profile question"
        else:
            typus = "Module"
        self.nodes.append(
            Node(
                title=node.title,
                typus=typus,
                depth=depth,
                number=".".join(number[:depth]),
                description=node.description,
                legal_reference=legal_reference,
                problem_description=problem_description,
            )
        )

        if IQuestionContainer.providedBy(node):
            i = 0
            depth += 1
            if len(number) < depth:
                number.append("0")
            for child in node.values():
                i += 1
                number[depth - 1] = str(i)
                self.AddToTree(child, depth, number)
        elif IRisk.providedBy(node):
            i = 0
            depth += 1
            number.append("0")
            for child in node.values():
                if not ISolution.providedBy(child):
                    continue
                i += 1
                number[depth - 1] = str(i)
                description = u"<ul>"
                for field in ("action_plan", "prevention_plan", "requirements"):
                    value = getattr(child, field, "") or ""
                    if value:
                        description = u"{0}<li>{1}</li>".format(description, value)
                description = u"{0}</ul>".format(description)
                self.nodes.append(
                    Node(
                        title=child.description,
                        typus="Measure",
                        depth=depth,
                        number=".".join(number[:depth]),
                        description=description,
                    )
                )

    def get_survey_nodes(self):
        i = 0
        for child in self.context.values():
            if IQuestionContainer.providedBy(child):
                i += 1
                self.AddToTree(child, depth=1, number=[str(i)])
        return self.nodes

    def get_data(self, for_download=False):
        """ Gets the data structure in a format suitable for `DocxCompiler`
        """
        title = self.context.aq_parent.title
        data = {
            "title": title,
            "heading": "",
            "section_headings": [title],
            "nodes": [self.get_survey_nodes()],
        }
        return data

    @property
    def _filename(self):
        """ Return the document filename
        """
        filename = _(
            "filename_tool_contents",
            default=u"Contents of OIRA tool ${title}",
            mapping=dict(title=self.context.title),
        )
        filename = translate(filename, context=self.request)
        return filename.encode("utf8") + ".docx"
